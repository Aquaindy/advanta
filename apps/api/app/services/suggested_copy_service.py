"""Suggested-copy persistence + file rendering.

Listing/fetching/deleting `SuggestedCopy` rows (always workspace-scoped) plus
rendering a copy — or a whole bundle — to a downloadable .txt or .docx file.

DOCX is produced with python-docx; bodies use light markdown (## headings,
- bullets) which we map to Word headings / list paragraphs so the export reads
cleanly in Word/Google Docs.
"""

from __future__ import annotations

from io import BytesIO
from uuid import UUID

from sqlalchemy.orm import Session

from app.models.suggested_copy import SuggestedCopy, SuggestedCopyType

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


# ---------------------------------------------------------------------------
# Queries
# ---------------------------------------------------------------------------


def list_suggested_copies(
    db: Session,
    *,
    workspace_id: UUID,
    copy_type: SuggestedCopyType | None = None,
    profile_id: UUID | None = None,
) -> list[SuggestedCopy]:
    query = db.query(SuggestedCopy).filter(
        SuggestedCopy.workspace_id == workspace_id
    )
    if copy_type is not None:
        query = query.filter(SuggestedCopy.copy_type == copy_type)
    if profile_id is not None:
        query = query.filter(SuggestedCopy.growth_dna_profile_id == profile_id)
    return query.order_by(SuggestedCopy.created_at.desc()).all()


def get_suggested_copy(
    db: Session, *, workspace_id: UUID, copy_id: UUID
) -> SuggestedCopy | None:
    return (
        db.query(SuggestedCopy)
        .filter(
            SuggestedCopy.id == copy_id,
            SuggestedCopy.workspace_id == workspace_id,
        )
        .first()
    )


def delete_suggested_copy(
    db: Session, *, workspace_id: UUID, copy_id: UUID
) -> bool:
    row = get_suggested_copy(db, workspace_id=workspace_id, copy_id=copy_id)
    if row is None:
        return False
    db.delete(row)
    db.commit()
    return True


# ---------------------------------------------------------------------------
# Filenames
# ---------------------------------------------------------------------------


def safe_filename(*parts: str) -> str:
    """Content-Disposition headers are latin-1 only — strip everything else."""
    raw = "_".join(p for p in parts if p).replace(" ", "_").replace("/", "-")
    cleaned = raw.encode("ascii", errors="ignore").decode("ascii")
    cleaned = "".join(ch for ch in cleaned if ch.isalnum() or ch in "._-")
    return (cleaned[:80] or "suggested_copy").strip("._-") or "suggested_copy"


# ---------------------------------------------------------------------------
# TXT rendering
# ---------------------------------------------------------------------------


def _txt_block(copy: SuggestedCopy) -> str:
    return (
        f"{copy.title}\n"
        f"{'=' * len(copy.title)}\n"
        f"Product / Service: {copy.product_name}\n"
        f"Section: {copy.section}\n"
        f"Type: {copy.copy_type.value}\n"
        f"\n{copy.body.strip()}\n"
    )


def render_txt(copy: SuggestedCopy) -> bytes:
    return _txt_block(copy).encode("utf-8")


def render_bundle_txt(copies: list[SuggestedCopy], *, product_name: str) -> bytes:
    header = (
        f"Suggested Copies — {product_name}\n"
        f"{'#' * 60}\n"
        f"{len(copies)} item(s)\n\n"
    )
    body = ("\n" + "-" * 60 + "\n\n").join(_txt_block(c) for c in copies)
    return (header + body).encode("utf-8")


# ---------------------------------------------------------------------------
# DOCX rendering (python-docx)
# ---------------------------------------------------------------------------


def _write_body_to_doc(doc, body: str) -> None:
    """Render a light-markdown body into a python-docx document."""
    for raw_line in body.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)


def _add_copy_to_doc(doc, copy: SuggestedCopy) -> None:
    doc.add_heading(copy.title, level=1)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Product / Service: {copy.product_name}  ·  "
        f"Section: {copy.section}  ·  Type: {copy.copy_type.value}"
    ).italic = True
    _write_body_to_doc(doc, copy.body)


def _new_document():
    from docx import Document  # imported lazily so a missing dep only hits DOCX

    return Document()


def render_docx(copy: SuggestedCopy) -> bytes:
    doc = _new_document()
    _add_copy_to_doc(doc, copy)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_bundle_docx(copies: list[SuggestedCopy], *, product_name: str) -> bytes:
    doc = _new_document()
    doc.add_heading(f"Suggested Copies — {product_name}", level=0)
    doc.add_paragraph(f"{len(copies)} item(s)")
    for idx, copy in enumerate(copies):
        if idx:
            doc.add_page_break()
        _add_copy_to_doc(doc, copy)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
